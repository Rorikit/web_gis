import io
from datetime import date, datetime

from django.db.models import Q
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet

from apps.accounts.models import District

from .models import AreaState, Damage, OrderKind

# Приложение №2: пункты 5-26 заполняются пользователем перед выводом отчёта
# (п. 1-4 берутся из базы, п. 15 — графическая схема, не заполняется).
DAMAGE_CARD_INPUT_FIELDS: list[tuple[str, str]] = [
    ('requestNumber', 'Заявка №'),
    ('requestDate', 'Дата заявки'),
    ('startChamber', '5. Начальная камера (точка) участка повреждения'),
    ('endChamber', '6. Конечная камера (точка) участка повреждения'),
    ('distanceFromChamber', '7. Расстояние до повреждения от начальной камеры, м'),
    ('layingType', '8. Тип прокладки тр-да в месте повреждения'),
    ('drainedPipe', '9. Дренируемый тр-д'),
    ('damagedPipe', '10. Поврежденный тр-д'),
    ('outerDiameter', '11. Наружный диаметр, мм'),
    ('mode', '12. Режим'),
    ('damagedElement', '13. Поврежденный элемент трубопровода'),
    ('damageNature', '14. Характер повреждения'),
    ('damageSizeA', '16. Размер повреждения, мм (A)'),
    ('damageSizeB', '16. Размер повреждения, мм (B)'),
    ('damageArea', '16. Площадь повреждения, кв.мм'),
    ('damageReason', '17. Причины повреждения'),
    ('responsiblePerson', '18. Ответственный за устранение повреждения (Ф.И.О., должность)'),
    ('pipeRepair', '19. Ремонт тр-да и элементов'),
    ('replacedLength', '20. Длина замененного уч-ка, м'),
    ('insulationRepair', '21. Ремонт изоляционной конструкции тр-да'),
    ('channelRepair', '22. Ремонт канала'),
    ('shutdownDate', '23. Отключено с (дата)'),
    ('shutdownTime', '23. Отключено с (время)'),
    ('restoreDate', '24. Включено в работу (дата)'),
    ('restoreTime', '24. Включено в работу (время)'),
    ('channelState', '25. Состояние конструкций канала (камеры)'),
    ('relatedReasons', '26. Сопутствующие причины'),
    ('notes', 'Примечания. Схемы. Пояснения'),
    ('filledByPosition', 'Карту заполнил: должность'),
    ('filledByName', 'Карту заполнил: Ф.И.О.'),
    ('filledAt', 'Карту заполнил: дата'),
]

NETWORK_TYPE_CODES = {'ОТ': '1 - отопление', 'ГВС': '2 - ГВС'}


def _blank(value) -> str:
    if value is None:
        return ''
    return str(value).strip()


def _fill(value, width: int = 20) -> str:
    """Значение пользователя либо линия для ручного заполнения."""
    text = _blank(value)
    return text if text else '_' * width


def _district_caption(name: str) -> str:
    """Название района в шапке: не дублируем слово «район», если оно уже есть."""
    text = _blank(name)
    if not text:
        return f'{"_" * 20} района'
    return text if 'район' in text.lower() else f'{text} района'


def _set_cell(cell, title: str, value: str = '', *, hint: str = '') -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.paragraphs[0].text = ''
    run = cell.paragraphs[0].add_run(title)
    run.bold = True
    run.font.size = Pt(9)

    if hint:
        hint_paragraph = cell.add_paragraph()
        hint_run = hint_paragraph.add_run(hint)
        hint_run.font.size = Pt(7)

    value_paragraph = cell.add_paragraph()
    value_run = value_paragraph.add_run(value)
    value_run.font.size = Pt(9)


def build_damage_card_document(damage: Damage, fields: dict | None = None, additional_info: str = '') -> bytes:
    """Карта повреждения по форме Приложения №2 к ТЗ."""
    values = dict(fields or {})
    district_name = damage.district.name if damage.district_id else ''

    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Pt(36)

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(
        f'КАРТА ПОВРЕЖДЕНИЯ № {_fill(damage.order_number or damage.id)}'
        f'          Заявка № {_fill(values.get("requestNumber"))}'
    )
    header_run.bold = True
    header_run.font.size = Pt(12)

    subheader = document.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader_run = subheader.add_run(
        f'{_district_caption(district_name)}          Дата заявки {_fill(values.get("requestDate"))}'
    )
    subheader_run.bold = True
    subheader_run.font.size = Pt(12)

    table = document.add_table(rows=0, cols=3)
    table.style = 'Table Grid'

    def add_row(cells: list[tuple[str, str, str]], span: bool = False):
        row = table.add_row()
        if span and len(cells) == 1:
            merged = row.cells[0].merge(row.cells[2])
            _set_cell(merged, cells[0][0], cells[0][1], hint=cells[0][2])
            return row
        for index, (title, value, hint) in enumerate(cells):
            _set_cell(row.cells[index], title, value, hint=hint)
        return row

    # п.1-п.4 — из базы
    add_row([
        ('1. Дата обнаружения повреждения:', _fill(damage.detected_at), ''),
        ('2. Источник теплоснабжения:', _fill(damage.heat_source), ''),
        ('3. Признак т/сети:', NETWORK_TYPE_CODES.get(damage.network_type, _fill(damage.network_type)),
         '1 - отопление; 2 - ГВС; 3 - теплоноситель'),
    ])
    add_row([
        ('4. Место повреждения (адрес):', _fill(damage.address, 30), ''),
        ('5. Начальная камера (точка) участка повреждения:', _fill(values.get('startChamber')), ''),
        ('6. Конечная камера (точка) участка повреждения:', _fill(values.get('endChamber')), ''),
    ])
    add_row([
        ('7. Расстояние до повреждения от начальной камеры:', f'{_fill(values.get("distanceFromChamber"), 12)} м', ''),
        ('8. Тип прокладки тр-да в месте повреждения:', _fill(values.get('layingType'), 6),
         '0 - камера; 1 - непроходной канал; 2 - полупроходной канал; 3 - в гильзе; 4 - по подвалу; 5 - надземная; 6 - в пределах ЦТП'),
        ('9. Дренируемый тр-д:', _fill(values.get('drainedPipe'), 6),
         '1 - подающий; 2 - обратный; 3 - оба тр-да'),
    ])
    add_row([
        ('10. Поврежденный тр-д:', _fill(values.get('damagedPipe'), 6), '1 - подающий; 2 - обратный'),
        ('11. Наружный диаметр:', f'{_fill(values.get("outerDiameter"), 12)} мм', ''),
        ('12. Режим:', _fill(values.get('mode'), 6), '1 - эксплуатация; 2 - опрессовка; 3 - др. испытания'),
    ])
    add_row([
        ('13. Поврежденный элемент трубопровода:', _fill(values.get('damagedElement'), 6),
         '1 - прямой уч-к; 2 - прямой уч-к в стене или Н.О.; 3 - прямой уч-к в скользящей опоре; 4 - отвод; '
         '5 - байпас; 6 - переход; 7 - сварной стык; 8 - задвижка; 9 - СК; 10 - СКУ; 11 - дренаж; 12 - воздушник; 13 - клапан'),
        ('14. Характер повреждения:', _fill(values.get('damageNature'), 6),
         '1 - разрыв; 2 - свищ; 3 - разгерметизация СК; 4 - механическая деформация трубы; 5 - дефект задвижки'),
        ('15. Место расположения центра повреждения на тр-де (указать стрелкой снаружи символа):',
         'прямой  (   )        (   )  обратный', ''),
    ])
    add_row([
        ('16. Размер повреждения:',
         f'{_fill(values.get("damageSizeA"), 10)} мм X {_fill(values.get("damageSizeB"), 10)} мм\n'
         f'площадь {_fill(values.get("damageArea"), 12)} кв.мм', ''),
        ('17. Причины повреждения:', _fill(values.get('damageReason'), 6),
         '1 - наружная коррозия; 2 - внутренняя коррозия; 3 - электрокоррозия; 4 - дефект металла; '
         '5 - превышение допуст. давления, гидроудар; 6 - дефект сварки; 7 - износ металла'),
        ('18. Ответственный за устранение повреждения (Ф.И.О., должность):', _fill(values.get('responsiblePerson'), 30), ''),
    ])

    section_row = table.add_row()
    merged = section_row.cells[0].merge(section_row.cells[2])
    merged.paragraphs[0].text = ''
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    section_run = merged.paragraphs[0].add_run('Описание ремонтных работ:')
    section_run.bold = True
    section_run.font.size = Pt(10)

    add_row([('19. Ремонт тр-да и элементов:', _fill(values.get('pipeRepair'), 6),
              '1 - поставлена заплата; 2 - заменен уч-к трубы; 3 - заварен свищ; 4 - поставлен хомут; 5 - уч-к отглушен; '
              '6 - заменен элемент (задвижка, СК и др.); 7 - элемент демонтирован; 8 - набит сальник; 9 - прочее; '
              '10 - ремонт элемента (задвижка, СК и др.)')], span=True)
    add_row([('20. (Для кода 2 из п.19) Длина замененного уч-ка:', f'{_fill(values.get("replacedLength"), 10)} м', '')], span=True)
    add_row([('21. Ремонт изоляционной конструкции тр-да:', _fill(values.get('insulationRepair'), 6),
              '1 - восстановлена полностью; 2 - выполнена только противокоррозионная покраска; 3 - работы не проводились')], span=True)
    add_row([
        ('22. Ремонт канала:', _fill(values.get('channelRepair'), 6),
         '1 - восстановлен старыми элементами; 2 - плиты перекрытия заменены на новые; 3 - конструкции канала заменены полностью'),
        ('23. Отключено с:', f'дата {_fill(values.get("shutdownDate"), 12)}  время {_fill(values.get("shutdownTime"), 10)}', ''),
        ('24. Включено в работу:', f'дата {_fill(values.get("restoreDate"), 12)}  время {_fill(values.get("restoreTime"), 10)}', ''),
    ])
    add_row([('25. Состояние конструкций канала (камеры):', _fill(values.get('channelState'), 6),
              '1 - конструкции целые; 2 - протекают швы; 3 - разрушено перекрытие; 4 - разрушены стенки; '
              '5 - разрушены металлоконструкции')], span=True)
    add_row([('26. Сопутствующие причины:', _fill(values.get('relatedReasons'), 12),
              '1 - постоянное или периодическое подтопление грунтовыми или дренажными (ливневыми) водами до контакта с тр-дом; '
              '2 - канал заилен до контакта с тр-дом; 3 - капёж воды на тр-д; '
              '4 - наличие коррозионных факторов в сетевой воде (кислород и др.); '
              '5 - развитие наружной коррозии в месте утечки сетевой воды; 6 - наличие блуждающих токов). '
              'Примечание: при наличии нескольких (до 3-х) факторов проставить коды по мере убывания степени их влияния')], span=True)

    notes_text = '\n'.join(part for part in [_blank(values.get('notes')), _blank(additional_info)] if part)
    add_row([('Примечания. Схемы. Пояснения', notes_text, '')], span=True)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Карту заполнил:').bold = True
    document.add_paragraph(
        f'Должность {_fill(values.get("filledByPosition"), 30)}          Ф.И.О. {_fill(values.get("filledByName"), 30)}'
    )
    document.add_paragraph(f'{_fill(values.get("filledAt"), 24)} г.          Подпись {"_" * 24}')

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _write_header(sheet: Worksheet, columns: list[str]) -> None:
    sheet.append(columns)
    for cell in sheet[sheet.max_row]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)


def _improvement_summary(item: Damage) -> str:
    parts = [
        ('Основная', item.improvement_main),
        ('Внутриквартальная дорога', item.improvement_inner_road),
        ('Тротуар', item.improvement_sidewalk),
        ('Отмостка', item.improvement_blind_area),
    ]
    selected = [label for label, flag in parts if flag]
    return ', '.join(selected) if selected else '-'


def build_current_table_workbook(items, entity_type: str) -> bytes:
    workbook = Workbook()
    sheet = workbook.active

    if entity_type == 'orders':
        # Приложение №5 «Ордера»
        sheet.title = 'Ордера'
        _write_header(sheet, [
            '№ п/п', 'Район', '№ ордера', 'Адрес', 'Текущ./Гарант.', 'Дата открытия ордера',
            'Ордер открыт до', 'Дата закрытия ордера', 'з/зона, м2', 'Асфальт, м2',
            'Основная/внутриквартальная дорога, тратуар, отмостка', 'бардюр/поребрик',
            'Состояние участка', 'Подрядчик/УРТС/Участок', '№ договора',
            'Дата подачи заявки для восстановления благоустройства',
            'Срок выполнения работ, согласно графика', 'Примечание', 'фотоотчет', 'геолакация на карте',
        ])
        for index, item in enumerate(items, start=1):
            point = getattr(item, 'gis_point', None)
            sheet.append([
                index,
                item.district.name if item.district_id else '-',
                item.order_number,
                item.address,
                item.order_kind or '',
                str(item.order_opened_at) if item.order_opened_at else '',
                str(item.order_valid_until) if item.order_valid_until else '',
                str(item.order_closed_at) if item.order_closed_at else '',
                item.green_zone_area,
                item.asphalt_area,
                _improvement_summary(item),
                item.curb_count,
                item.area_state,
                item.contractor_type,
                item.contract_number,
                str(item.contractor_request_date) if item.contractor_request_date else '',
                str(item.planned_finish_date) if item.planned_finish_date else '',
                item.note,
                f'{len(item.photos.all())} фото',
                f'{point.latitude}, {point.longitude}' if point else '',
            ])
    else:
        # Приложение №3 «Общая»
        sheet.title = 'Повреждения'
        _write_header(sheet, [
            '№ п/п', 'Район', 'Адрес', 'ОТ/ГВС', 'Дата обнаружения', 'Дата устранения', '№ ордера',
            'Дата открытия ордера', 'Ордер открыт до', 'От какого источника запитан',
            'Текущее/Гидравлическое', 'Адреса отключенных абонентов', 'Характер повреждения',
            'Текущ./Гарант.', 'з/зона, м2', 'Асфальт, м2', 'Основная', 'Внутриквартальная дорога',
            'Тротуар', 'Отмостка', 'бардюр/поребрик', 'Состояние участка', 'Подрядчик/УРТС/Участок',
            '№ договора', 'Дата подачи заявки для восстановления благоустройства',
            'Срок выполнения работ, согласно графика', 'Примечание', 'Дата закрытия ордера',
            'фотоотчет', 'геолакация на карте',
        ])
        for index, item in enumerate(items, start=1):
            point = getattr(item, 'gis_point', None)
            sheet.append([
                index,
                item.district.name if item.district_id else '-',
                item.address,
                item.network_type,
                str(item.detected_at) if item.detected_at else '',
                str(item.fixed_at) if item.fixed_at else '',
                item.order_number,
                str(item.order_opened_at) if item.order_opened_at else '',
                str(item.order_valid_until) if item.order_valid_until else '',
                item.heat_source,
                item.damage_type,
                item.disconnected_addresses,
                item.damage_description,
                item.order_kind or '',
                item.green_zone_area,
                item.asphalt_area,
                'Да' if item.improvement_main else 'Нет',
                'Да' if item.improvement_inner_road else 'Нет',
                'Да' if item.improvement_sidewalk else 'Нет',
                'Да' if item.improvement_blind_area else 'Нет',
                item.curb_count,
                item.area_state,
                item.contractor_type,
                item.contract_number,
                str(item.contractor_request_date) if item.contractor_request_date else '',
                str(item.planned_finish_date) if item.planned_finish_date else '',
                item.note,
                str(item.order_closed_at) if item.order_closed_at else '',
                f'{len(item.photos.all())} фото',
                f'{point.latitude}, {point.longitude}' if point else '',
            ])

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _parse_report_date(report_date: str) -> date:
    try:
        return datetime.strptime(report_date, '%Y-%m-%d').date()
    except (TypeError, ValueError):
        return date.today()


def _reference_counters(report_day: date) -> dict[str, list[int]]:
    """Считает показатели Приложения №1 по каждому району."""
    year_start = date(report_day.year, 1, 1)
    still_open = Q(order_closed_at__isnull=True) | Q(order_closed_at__gt=report_day)
    opened_before = Q(order_opened_at__isnull=False, order_opened_at__lt=report_day)

    counters: dict[str, list[int]] = {}
    queryset = Damage.objects.filter(order_opened_at__isnull=False)

    def bump(district_id: str, index: int) -> None:
        counters.setdefault(district_id, [0, 0, 0, 0, 0, 0])[index] += 1

    closed_in_year = queryset.filter(order_closed_at__gte=year_start, order_closed_at__lte=report_day)
    for district_id, order_kind in closed_in_year.values_list('district_id', 'order_kind'):
        bump(district_id, 0 if order_kind == OrderKind.CURRENT else 1)

    open_on_date = queryset.filter(opened_before).filter(still_open)
    for district_id, order_kind, area_state in open_on_date.values_list('district_id', 'order_kind', 'area_state'):
        bump(district_id, 2 if order_kind == OrderKind.CURRENT else 3)
        if area_state == AreaState.IN_PROGRESS:
            bump(district_id, 4)
        elif area_state == AreaState.READY_TO_CLOSE:
            bump(district_id, 5)

    return counters


def build_reference_workbook(report_date: str) -> bytes:
    """Справка по форме Приложения №1 к ТЗ."""
    report_day = _parse_report_date(report_date)
    counters = _reference_counters(report_day)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'Справка'

    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    sheet.merge_cells('A1:G1')
    sheet['A1'] = 'СПРАВКА'
    sheet['A1'].font = Font(bold=True, size=12)
    sheet['A1'].alignment = center

    sheet.merge_cells('A2:G2')
    sheet['A2'] = (
        'о ходе работ по асфальтированию и благоустройству участков тепловых сетей '
        f'после аварийно-восстановительных работ на {report_day:%d.%m.%Y}'
    )
    sheet['A2'].alignment = center

    header_top = 4
    header_bottom = 5
    sheet.merge_cells(start_row=header_top, start_column=1, end_row=header_bottom, end_column=1)
    sheet.cell(row=header_top, column=1, value='Район')
    sheet.merge_cells(start_row=header_top, start_column=2, end_row=header_top, end_column=3)
    sheet.cell(row=header_top, column=2, value=f'Закрыто в {report_day.year} году')
    sheet.merge_cells(start_row=header_top, start_column=4, end_row=header_top, end_column=5)
    sheet.cell(row=header_top, column=4, value='Открыто')
    sheet.merge_cells(start_row=header_top, start_column=6, end_row=header_top, end_column=7)
    sheet.cell(row=header_top, column=6, value=f'в т.ч. информация по работам на {report_day:%d.%m.%Y}')

    for column, title in enumerate(['Текущие', 'Гарантийные', 'Текущие', 'Гарантийные', 'в работе', 'готов к закрытию'], start=2):
        sheet.cell(row=header_bottom, column=column, value=title)

    for row in (header_top, header_bottom):
        for column in range(1, 8):
            cell = sheet.cell(row=row, column=column)
            cell.font = Font(bold=True)
            cell.alignment = center
            cell.border = border

    totals = [0, 0, 0, 0, 0, 0]
    row_index = header_bottom + 1
    for district in District.objects.all():
        values = counters.get(district.id, [0, 0, 0, 0, 0, 0])
        sheet.cell(row=row_index, column=1, value=district.name).border = border
        for offset, value in enumerate(values):
            totals[offset] += value
            cell = sheet.cell(row=row_index, column=2 + offset, value=value)
            cell.alignment = center
            cell.border = border
        row_index += 1

    total_row = row_index
    sheet.cell(row=total_row, column=1, value='Итого').font = Font(bold=True)
    sheet.cell(row=total_row, column=1).border = border
    for offset, value in enumerate(totals):
        cell = sheet.cell(row=total_row, column=2 + offset, value=value)
        cell.font = Font(bold=True)
        cell.alignment = center
        cell.border = border

    grand_row = total_row + 1
    sheet.cell(row=grand_row, column=1, value='ВСЕГО').font = Font(bold=True)
    sheet.cell(row=grand_row, column=1).border = border
    sheet.merge_cells(start_row=grand_row, start_column=2, end_row=grand_row, end_column=3)
    sheet.cell(row=grand_row, column=2, value=totals[0] + totals[1])
    sheet.merge_cells(start_row=grand_row, start_column=4, end_row=grand_row, end_column=5)
    sheet.cell(row=grand_row, column=4, value=totals[2] + totals[3])
    sheet.cell(row=grand_row, column=6, value=totals[4])
    sheet.cell(row=grand_row, column=7, value=totals[5])
    for column in range(2, 8):
        cell = sheet.cell(row=grand_row, column=column)
        cell.font = Font(bold=True)
        cell.alignment = center
        cell.border = border

    sheet.column_dimensions['A'].width = 26
    for column in range(2, 8):
        sheet.column_dimensions[get_column_letter(column)].width = 16

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
